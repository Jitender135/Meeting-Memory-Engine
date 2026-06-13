"""
ingest.py — Multi-format transcript ingestion pipeline.

Supported formats:
    .txt  — plain text transcripts
    .pdf  — Zoom/Google Meet PDF exports (via PyMuPDF)
    .docx — Word meeting notes (via python-docx)

Design decision — why format-specific loaders over LangChain's generic ones:
    LangChain's generic loaders lose metadata and formatting context.
    Format-specific extraction gives us clean text we can reliably
    parse for meeting title, date, and participants.
"""

import fitz  # PyMuPDF
import chromadb

from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
# HuggingFace replaced by Jina API embeddings for deployment compatibility

# ── Environment ───────────────────────────────────────────────────
load_dotenv(Path(__file__).resolve().parents[2] / ".env")

# ── Constants ─────────────────────────────────────────────────────
CHROMA_PATH   = Path(__file__).resolve().parents[2] / "chroma_db"
DATA_PATH     = Path(__file__).resolve().parents[2] / "data"
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 50
SUPPORTED     = {".txt", ".pdf", ".docx"}


# ─────────────────────────────────────────────────────────────────
# FORMAT-SPECIFIC TEXT EXTRACTORS
# ─────────────────────────────────────────────────────────────────

def extract_txt(file_path: Path) -> str:
    """Read plain text transcript."""
    return file_path.read_text(encoding="utf-8")


def extract_pdf(file_path: Path) -> str:
    """
    Extract text from PDF using PyMuPDF.
    Why PyMuPDF over pdfplumber or PyPDF2:
      - Fastest parser available in Python
      - Preserves reading order better than alternatives
      - Handles scanned PDFs with embedded text layers
    """
    doc   = fitz.open(str(file_path))
    pages = [page.get_text("text") for page in doc]
    doc.close()
    return "\n".join(pages).strip()


def extract_docx(file_path: Path) -> str:
    """
    Extract text from DOCX paragraph by paragraph.
    Why not use LangChain's Docx2txtLoader:
      - Docx2txt strips paragraph boundaries, losing structure
      - python-docx preserves paragraph order and empty lines,
        which our chunker uses as semantic split points
    """
    doc        = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs).strip()


# ─────────────────────────────────────────────────────────────────
# METADATA EXTRACTION
# ─────────────────────────────────────────────────────────────────

def extract_metadata(file_path: Path) -> dict:
    """
    Pull structured metadata from filename and file content.

    Primary strategy  — parse date from filename: meeting_YYYY_MM_DD.*
    Fallback strategy — timestamp 0 (chunk still ingested, just unfiltered)

    Why store as both string and integer:
      - meeting_date (string) → displayed in UI citations
      - meeting_timestamp (int) → used for ChromaDB $gte/$lte filtering
        because ChromaDB where clauses only support numeric comparisons
    """
    stem  = file_path.stem   # e.g. meeting_2024_08_15
    parts = stem.split("_")

    try:
        date_str  = f"{parts[1]}-{parts[2]}-{parts[3]}"
        date_obj  = datetime.strptime(date_str, "%Y-%m-%d")
        timestamp = int(date_obj.timestamp())
    except (IndexError, ValueError):
        date_str  = "unknown"
        timestamp = 0

    return {
        "meeting_date":      date_str,
        "meeting_timestamp": timestamp,
        "source_file":       file_path.name,
        "file_format":       file_path.suffix.lstrip("."),
    }


def extract_title(raw_text: str) -> str:
    """
    Pull meeting title from the first non-empty line of the transcript.
    Works across TXT, PDF, and DOCX since all follow the same header format.
    """
    for line in raw_text.splitlines():
        line = line.strip()
        if line:
            return line.replace("Meeting Title:", "").strip()
    return "Untitled Meeting"


# ─────────────────────────────────────────────────────────────────
# DOCUMENT LOADING
# ─────────────────────────────────────────────────────────────────

def load_file(file_path: Path) -> Document | None:
    """
    Load a single file into a LangChain Document object.
    Returns None if the format is unsupported or extraction fails.
    """
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED:
        print(f"  Skipping unsupported format: {file_path.name}")
        return None

    try:
        if suffix == ".txt":
            raw_text = extract_txt(file_path)
        elif suffix == ".pdf":
            raw_text = extract_pdf(file_path)
        elif suffix == ".docx":
            raw_text = extract_docx(file_path)

        if not raw_text.strip():
            print(f"  Skipping empty file: {file_path.name}")
            return None

        metadata          = extract_metadata(file_path)
        metadata["title"] = extract_title(raw_text)

        return Document(page_content=raw_text, metadata=metadata)

    except Exception as e:
        print(f"  Error loading {file_path.name}: {e}")
        return None


# ─────────────────────────────────────────────────────────────────
# CHUNKING
# ─────────────────────────────────────────────────────────────────

def chunk_documents(docs: list[Document]) -> list[Document]:
    """
    Split documents into chunks preserving metadata on every chunk.

    Why RecursiveCharacterTextSplitter with these separators:
      Tries splits in order: paragraph → newline → sentence → word
      This preserves semantic boundaries — a decision or action item
      stays in one chunk rather than being split mid-sentence.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " "],
    )
    return splitter.split_documents(docs)


# ─────────────────────────────────────────────────────────────────
# EMBEDDING
# ─────────────────────────────────────────────────────────────────

def get_embedding_function():
    """
    Jina AI embeddings — API-based, no local model download.
    Why Jina over local HuggingFace for deployment:
        Local HuggingFace models require 400-500MB RAM to load.
        Free tier servers cap at 512MB total — causes OOM crash.
        Jina API uses <10MB RAM, deployable on any free tier.
        jina-embeddings-v2-base-en outperforms all-MiniLM-L6-v2
        on retrieval benchmarks.
    """
    import os
    from langchain_community.embeddings import JinaEmbeddings
    return JinaEmbeddings(
        jina_api_key=os.getenv("JINA_API_KEY"),
        model_name="jina-embeddings-v2-base-en",
    )


# ─────────────────────────────────────────────────────────────────
# MAIN INGEST FUNCTION
# ─────────────────────────────────────────────────────────────────

def ingest_all(data_path: Path = DATA_PATH) -> dict:
    """
    Full ingest pipeline:
      1. Discover all supported files in data/
      2. Extract text per format (TXT / PDF / DOCX)
      3. Attach metadata (date, title, format, source file)
      4. Chunk with overlap
      5. Embed with HuggingFace all-MiniLM-L6-v2
      6. Upsert into ChromaDB (safe to re-run — no duplicates)

    Returns a summary dict consumed by FastAPI's /ingest endpoint.
    """
    # ── Discover files ────────────────────────────────────────────
    all_files = [
        f for f in data_path.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED
    ]

    if not all_files:
        return {
            "status":  "error",
            "message": f"No supported files found in {data_path}. "
                       f"Supported formats: {', '.join(SUPPORTED)}",
        }

    print(f"Found {len(all_files)} file(s). Starting ingest...")

    # ── Load all files ────────────────────────────────────────────
    all_docs = []
    for f in sorted(all_files):
        doc = load_file(f)
        if doc:
            all_docs.append(doc)
            fmt = f.suffix.upper().lstrip(".")
            print(f"  [{fmt}] Loaded: {f.name} — {doc.metadata['title']}")

    if not all_docs:
        return {"status": "error", "message": "All files failed to load."}

    # ── Chunk ─────────────────────────────────────────────────────
    chunks = chunk_documents(all_docs)
    print(f"  Total chunks: {len(chunks)}")

    # ── Embed + store ─────────────────────────────────────────────
    client     = chromadb.PersistentClient(path=str(CHROMA_PATH))
    collection = client.get_or_create_collection(
        name="meeting_transcripts",
        metadata={"hnsw:space": "cosine"},
    )

    embed_fn = get_embedding_function()

    for i, chunk in enumerate(chunks):
        embedding = embed_fn.embed_query(chunk.page_content)
        collection.upsert(
            ids=[f"chunk_{i}"],
            embeddings=[embedding],
            documents=[chunk.page_content],
            metadatas=[chunk.metadata],
        )

    print(f"  Ingested {len(chunks)} chunks into ChromaDB.")

    return {
        "status":     "success",
        "files":      len(all_docs),
        "chunks":     len(chunks),
        "collection": "meeting_transcripts",
        "formats":    list({f.suffix.lstrip(".") for f in all_files}),
    }


# ── Run directly for testing ──────────────────────────────────────
if __name__ == "__main__":
    result = ingest_all()
    print("\nIngest complete:", result)